
from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

import docx  
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

_MARGEM_PADRAO_PT = 72  


_FONTE_PADRAO_TTF = Path(__file__).resolve().parent.parent / "fontes" / "Renaissance-Initialen 1.ttf"




def registrar_fonte(caminho_ttf: Path) -> str:

    caminho_ttf = Path(caminho_ttf)
    if not caminho_ttf.exists():
        raise FileNotFoundError(f"fonte não encontrada: {caminho_ttf}")
    nome = caminho_ttf.stem.replace(" ", "_")
    pdfmetrics.registerFont(TTFont(nome, str(caminho_ttf)))
    return nome


def _resolver_nome_fonte(fonte_ttf: Optional[Path]) -> Optional[str]:

    if fonte_ttf:
        return registrar_fonte(Path(fonte_ttf))
    if _FONTE_PADRAO_TTF.exists():
        return registrar_fonte(_FONTE_PADRAO_TTF)
    return None




def _escapar(texto: str) -> str:
    return texto.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


_MAPA_ALINHAMENTO = {
    0: TA_LEFT,     
    1: TA_CENTER,   
    2: TA_RIGHT,    
    3: TA_JUSTIFY,  
    None: TA_LEFT,
}


def _ler_docx(caminho: Path):

    d = docx.Document(str(caminho))
    secao = d.sections[0]
    margens = {
        "top": secao.top_margin.pt if secao.top_margin else _MARGEM_PADRAO_PT,
        "bottom": secao.bottom_margin.pt if secao.bottom_margin else _MARGEM_PADRAO_PT,
        "left": secao.left_margin.pt if secao.left_margin else _MARGEM_PADRAO_PT,
        "right": secao.right_margin.pt if secao.right_margin else _MARGEM_PADRAO_PT,
    }
    tamanho_pagina = (
        secao.page_width.pt if secao.page_width else A4[0],
        secao.page_height.pt if secao.page_height else A4[1],
    )

    blocos: list[tuple[str, int | None]] = []
    for p in d.paragraphs:
        alinhamento = p.alignment.value if p.alignment is not None else None
        blocos.append((p.text, alinhamento))
    for tabela in d.tables:
        for linha in tabela.rows:
            texto_linha = " | ".join(celula.text for celula in linha.cells)
            blocos.append((texto_linha, None))
    return blocos, margens, tamanho_pagina


def _rtf_para_texto(caminho: Path) -> str:
    from striprtf.striprtf import rtf_to_text
    conteudo = caminho.read_text(encoding="utf-8", errors="ignore")
    return rtf_to_text(conteudo)




def _gerar_pdf(blocos: list[tuple[str, int | None]], margens: dict,
                tamanho_pagina: tuple[float, float], nome_fonte: Optional[str]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=tamanho_pagina,
        topMargin=margens["top"], bottomMargin=margens["bottom"],
        leftMargin=margens["left"], rightMargin=margens["right"],
    )

    fonte = nome_fonte or "Helvetica"
    estilo_base = ParagraphStyle("base", fontName=fonte, fontSize=12, leading=16)

    elementos = []
    for texto, alinhamento in blocos:
        if not texto.strip():
            elementos.append(Spacer(1, 10))
            continue
        estilo = ParagraphStyle(
            "p", parent=estilo_base,
            alignment=_MAPA_ALINHAMENTO.get(alinhamento, TA_LEFT),
        )
        elementos.append(Paragraph(_escapar(texto), estilo))
        elementos.append(Spacer(1, 6))

    if not elementos:
        elementos = [Paragraph("", estilo_base)]

    doc.build(elementos)
    return buf.getvalue()


def converter_para_pdf(caminho_entrada: str | Path, fonte_ttf: Optional[Path] = None) -> bytes:

    caminho_entrada = Path(caminho_entrada)
    extensao = caminho_entrada.suffix.lower()
    if extensao not in (".docx", ".rtf", ".txt"):
        raise ValueError(f"formato não suportado: {extensao} (use .docx, .rtf ou .txt)")

    nome_fonte = _resolver_nome_fonte(Path(fonte_ttf) if fonte_ttf else None)

    if extensao == ".docx":
        blocos, margens, tamanho_pagina = _ler_docx(caminho_entrada)
    else:
        if extensao == ".rtf":
            texto = _rtf_para_texto(caminho_entrada)
        else:  # .txt
            texto = caminho_entrada.read_text(encoding="utf-8", errors="ignore")
        blocos = [(linha, None) for linha in texto.splitlines()] or [("", None)]
        margens = {"top": _MARGEM_PADRAO_PT, "bottom": _MARGEM_PADRAO_PT,
                   "left": _MARGEM_PADRAO_PT, "right": _MARGEM_PADRAO_PT}
        tamanho_pagina = A4

    return _gerar_pdf(blocos, margens, tamanho_pagina, nome_fonte)
